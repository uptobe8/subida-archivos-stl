from pathlib import Path
import html
import json
import re
import zipfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent / "output"
OUT.mkdir(parents=True, exist_ok=True)
V = "4.27.4"
NAVY = "#031336"
BLUE = "#0B2B6B"
PURPLE = "#261C30"
VIOLET = "#08007C"
GOLD = "#E09900"
WHITE = "#FFFFFF"
SOFT = "#F4F6FA"
INK = "#243467"
MUTED = "#59647A"
HEAD = "Josefin Sans"
BODY = "Josefin Sans"
KICK = "Catamaran"
HOME_IMG = "http://185.232.41.199/~wwwlegadianew/wp-content/uploads/2026/02/IMG_6289.jpg"
IMG = {
    "hero": "https://images.pexels.com/photos/8112164/pexels-photo-8112164.jpeg?auto=compress&dpr=1&h=1000&w=1600",
    "search": "https://images.pexels.com/photos/7875863/pexels-photo-7875863.jpeg?auto=compress&dpr=1&h=1000&w=1600",
    "process": "https://images.pexels.com/photos/4427422/pexels-photo-4427422.jpeg?auto=compress&dpr=1&h=1000&w=1600",
    "rights": "https://images.pexels.com/photos/8112153/pexels-photo-8112153.jpeg?auto=compress&dpr=1&h=1000&w=1600",
    "team": "https://images.pexels.com/photos/7841445/pexels-photo-7841445.jpeg?auto=compress&dpr=1&h=1000&w=1600",
    "docs": "https://images.pexels.com/photos/7841846/pexels-photo-7841846.jpeg?auto=compress&dpr=1&h=1000&w=1600",
    "talk": "https://images.pexels.com/photos/7876093/pexels-photo-7876093.jpeg?auto=compress&dpr=1&h=1000&w=1600",
    "scene": "https://images.pexels.com/photos/4427626/pexels-photo-4427626.jpeg?auto=compress&dpr=1&h=1200&w=2000",
    "contacted": "http://185.232.41.199/~wwwlegadianew/wp-content/uploads/2025/12/business-consulting-5-sdc-16.jpg",
    "partners": "http://185.232.41.199/~wwwlegadianew/wp-content/uploads/2025/12/business-consulting-5-sdc-6.jpg",
}

TEXT = {
    "hero_k": "Soluciones Legadia",
    "hero_h": "Servicios especializados para desbloquear herencias",
    "hero_p": "Investigamos, ordenamos y tramitamos expedientes hereditarios cuando hay herederos desconocidos, documentación incompleta, derechos hereditarios pendientes o comunicaciones que requieren una respuesta profesional.",
    "hero_label": "Investigación · Tramitación · Valoración",
    "hero_card": "Por qué conviene clasificar el caso antes de actuar",
    "services_k": "Tres vías de actuación",
    "services_h": "Elige la vía que encaja con el bloqueo.",
    "services_p": "Búsqueda, tramitación o valoración de derechos: cada landing mantiene una estructura específica para resolver dudas, cualificar el caso y facilitar el contacto.",
    "s1_h": "Búsqueda de Herederos",
    "s1_p": "Pasar de “no sabemos quién debe actuar” a una línea de investigación y contacto clara.",
    "s2_h": "Tramitación de Herencias",
    "s2_p": "Un expediente hereditario necesita método, calendario, documentos correctos y comunicación entre partes.",
    "s3_h": "Compra de Derechos Hereditarios",
    "s3_p": "Convertir incertidumbre, espera y bloqueo en una alternativa clara de valoración y decisión.",
    "access_k": "Otros accesos Legadia",
    "access_h": "También podemos ayudarte desde aquí.",
    "access_p": "Si hemos contactado contigo o representas a una entidad profesional, accede directamente al área correspondiente.",
    "contacted_h": "Te hemos contactado",
    "contacted_p": "Si ha recibido una llamada o una comunicación por nuestra parte, no ha sido al azar. Investigamos, verificamos y localizamos herederos legítimos para ayudarles a comprender su caso y recuperar lo que les puede corresponder.",
    "partners_h": "Colaboraciones",
    "partners_p": "Colaboramos con notarías, asesorías y otros profesionales cuando un expediente necesita investigación, verificación y localización de herederos con un método claro y coordinado.",
    "diag_k": "Diagnóstico inicial",
    "diag_h": "Qué tipo de bloqueo necesitas resolver",
    "diag_p": "Cada caso hereditario tiene una puerta de entrada distinta. Puede comenzar por una búsqueda de herederos, por una tramitación paralizada, por la necesidad de vender derechos hereditarios o por una comunicación inesperada.",
    "diag": [
        "Necesitas localizar a herederos antes de poder cerrar una operación o expediente.",
        "La herencia está iniciada, pero nadie consigue ordenar documentación, impuestos y adjudicación.",
        "Tienes derechos hereditarios y quieres valorar una salida rápida, segura y documentada.",
        "Has recibido una comunicación de Legadia y necesitas entender el motivo y tus opciones.",
    ],
    "gallery_k": "Investigación aplicada",
    "gallery_h": "Investigación, documentación y contacto con enfoque profesional.",
    "gallery_p": "Imágenes reales de Legadia y recursos visuales de apoyo para que la página transmita equipo, proceso documental y contexto de Madrid desde el primer recorrido.",
    "gallery": [
        "Especialistas coordinando expedientes hereditarios.",
        "Revisión de datos, registros y pruebas familiares.",
        "Comunicación clara para personas afectadas.",
    ],
    "class_k": "Antes de elegir servicio",
    "class_h": "Por qué conviene clasificar el caso antes de actuar",
    "class_p1": "Una herencia bloqueada no se resuelve siempre con la misma vía. A veces el problema es documental, otras veces genealógico, otras jurídico y, en ocasiones, económico. Actuar sin clasificar el expediente puede multiplicar gestiones y retrasar la solución.",
    "class_p2": "La especialización permite decidir si procede investigar herederos, tramitar la herencia, valorar derechos hereditarios o abrir una comunicación formal con posibles interesados.",
    "class_q": "El objetivo no es hacer más trámites, sino saber cuál es el trámite correcto.",
    "method_k": "Proceso común",
    "method_h": "Metodología de actuación Legadia.",
    "method_p": "Trabajamos cada servicio con una lógica común: diagnóstico, documentación, investigación, propuesta de actuación y comunicación clara del siguiente paso.",
    "method": [
        ("01 · Diagnóstico inicial", "Identificamos el origen del bloqueo y a quién afecta."),
        ("02 · Investigación y contraste", "Revisamos documentos, registros, datos familiares y bienes relacionados."),
        ("03 · Plan de actuación", "Definimos la vía adecuada: búsqueda, tramitación, valoración o contacto."),
        ("04 · Acompañamiento", "Acompañamos la ejecución hasta que el caso pueda avanzar con seguridad."),
    ],
    "aud_k": "Para quién",
    "aud_h": "Para quién está pensada esta valoración.",
    "aud_p": "Si te reconoces en alguna de estas situaciones, conviene revisar el expediente antes de seguir acumulando gestiones, costes o comunicaciones sin respuesta clara.",
    "aud": [
        "Quien tiene un inmueble o expediente sin titular claro.",
        "Quien necesita cerrar una herencia y no sabe por dónde empezar.",
        "Quien quiere saber si una comunicación recibida es legítima y qué implica.",
        "Quien busca una alternativa para vender o valorar derechos hereditarios.",
    ],
    "err_k": "Qué evitamos",
    "err_h": "Errores que pueden alargar el bloqueo.",
    "err_p": "La intervención temprana no consiste en hacer más gestiones, sino en evitar pasos que consumen tiempo y no acercan el expediente a una solución.",
    "err_sub": "Antes de avanzar, revisamos especialmente:",
    "err": [
        "Elegir un servicio solo por el nombre sin revisar el origen del bloqueo.",
        "Duplicar gestiones entre varias personas sin coordinar documentación.",
        "Dejar pasar plazos por no saber quién debe actuar.",
        "Pedir documentos de forma indiscriminada sin clasificar el caso.",
    ],
    "scen_k": "Escenarios frecuentes",
    "scen_h": "Qué servicio suele encajar en cada caso",
    "scen_p": "Estos casos explican por qué conviene ordenar la información antes de tomar decisiones, firmar documentos o dejar pasar más tiempo.",
    "scen": [
        "Si falta identificar a los sucesores, la vía suele empezar por búsqueda de herederos.",
        "Si los herederos existen pero el trámite está parado, se revisa la tramitación documental.",
        "Si el problema es económico o de salida rápida, puede estudiarse la compra o valoración de derechos.",
        "Si has recibido una comunicación, primero conviene verificar el motivo y el alcance del expediente.",
    ],
    "doc_k": "Documentación mínima",
    "doc_h": "Qué conviene tener preparado para avanzar.",
    "doc_p": "No hace falta tenerlo todo resuelto antes de contactar. Estos datos ayudan a clasificar el expediente y pedir solo la información necesaria.",
    "docs": [
        "Datos de la persona fallecida y relación conocida con el expediente.",
        "Documentos sucesorios disponibles: testamento, certificados, escrituras, notas simples o comunicaciones.",
        "Información sobre bienes, cargas, deudas, comunidades o terceros afectados.",
        "Objetivo del contacto: localizar, tramitar, vender, responder o desbloquear.",
    ],
    "crit_k": "Criterio profesional",
    "crit_h": "El objetivo no es hacer más trámites, sino saber cuál es el trámite correcto.",
    "crit_p1": "Cuando una herencia está bloqueada, actuar sin diagnóstico puede aumentar la confusión. Por eso cada landing está enfocada en reconocer el problema, explicar qué puede haber detrás y facilitar una consulta clara.",
    "crit_p2": "El siguiente paso debe ser sencillo: contar el caso, aportar lo mínimo necesario y recibir una orientación sobre la vía razonable.",
    "trust_k": "Confianza",
    "trust_h": "Experiencia, método y trazabilidad.",
    "trust_p": "En expedientes hereditarios sensibles, la confianza se construye explicando qué se hace, por qué se hace y qué documentación respalda cada paso.",
    "trust_card_h": "Casos explicados de forma clara",
    "trust_card_p": "Recursos visuales para entender situaciones frecuentes sin exponer datos personales ni expedientes confidenciales.",
    "form_k": "Consulta inicial",
    "form_h": "Cuéntanos qué ocurre y te orientamos sobre el siguiente paso.",
    "form_p": "Cuéntanos qué ocurre con la herencia, qué documentación tienes y qué necesitas resolver. Revisaremos la información para orientarte sobre el siguiente paso razonable.",
    "direct_h": "Contacto directo",
    "direct_p": "También puedes contactar por teléfono, email o WhatsApp para explicar brevemente el caso.",
    "faq_k": "Preguntas frecuentes",
    "faq_h": "Dudas antes de contactar.",
    "faq_p": "Preguntas reales que suelen frenar la consulta cuando una herencia está bloqueada, hay herederos desconocidos o falta documentación.",
    "faq": [
        ("¿Qué servicio debo elegir?", "Si no lo tienes claro, puedes enviar el caso como consulta general. Legadia clasifica la vía más adecuada tras revisar el origen del bloqueo."),
        ("¿Puedo consultar aunque tenga poca documentación?", "Sí. Con datos básicos puede iniciarse una primera orientación y solicitar solo la documentación realmente necesaria."),
        ("¿Trabajáis solo para herederos?", "No necesariamente. También pueden existir propietarios, comunidades, administradores, entidades o terceros afectados por una herencia yacente."),
        ("¿Se puede combinar búsqueda y tramitación?", "Sí. Muchos expedientes empiezan localizando herederos y continúan con una fase documental o de tramitación sucesoria."),
    ],
    "final_k": "Primer paso",
    "final_h": "No tienes que saber qué servicio elegir antes de contactar.",
    "final_p": "Explica brevemente qué ocurre. Revisaremos el origen del bloqueo para indicarte la vía razonable y la documentación realmente necesaria.",
}


def attrs(values):
    return "".join(f' {k}="{html.escape(str(v), quote=True)}"' for k, v in values.items() if v is not None)


def sc(name, values=None, body=""):
    return f"[{name}{attrs(values or {})}]{body}[/{name}]"


def base():
    return {"_builder_version": V, "_module_preset": "default", "global_colors_info": "{}"}


def section(body, label, **extra):
    a = base() | {"fb_built": "1", "admin_label": label} | extra
    return sc("et_pb_section", a, body)


def row(body, structure="1_1", **extra):
    a = base() | {"column_structure": structure, "width": "86%", "max_width": "1240px", "width_tablet": "92%", "width_phone": "calc(100% - 32px)", "custom_padding": "0px||0px||false|false"} | extra
    return sc("et_pb_row", a, body)


def col(body, kind="1_1", **extra):
    return sc("et_pb_column", base() | {"type": kind} | extra, body)


def blurb(title, body="", *, level="h3", dark=False, bg="RGBA(255,255,255,0)", pad="0px", icon=None, title_size="28px", body_size="17px", align="left", label=None, **extra):
    a = base() | {
        "title": title,
        "admin_label": label or f"Componente nativo · {title}",
        "header_level": level,
        "use_icon": "on" if icon else "off",
        "font_icon": f"{icon}||divi||400" if icon else None,
        "icon_color": GOLD,
        "use_icon_font_size": "on" if icon else "off",
        "icon_font_size": "32px",
        "icon_placement": "top",
        "text_orientation": align,
        "header_font": f"{HEAD}|300|||||||",
        "header_text_color": WHITE if dark else BLUE,
        "header_font_size": title_size,
        "header_font_size_tablet": title_size,
        "header_font_size_phone": "24px" if int(re.sub(r'\D','',title_size) or 28) > 30 else title_size,
        "header_line_height": "1.16em",
        "body_font": f"{BODY}|300|||||||",
        "body_text_color": "rgba(255,255,255,0.80)" if dark else MUTED,
        "body_font_size": body_size,
        "body_font_size_phone": "16px",
        "body_line_height": "1.58em",
        "background_color": bg,
        "custom_padding": pad,
        "animation_style": "fade",
        "animation_direction": "bottom",
        "animation_duration": "650ms",
    } | extra
    content = f"<p>{html.escape(body)}</p>" if body else ""
    return sc("et_pb_blurb", a, content)


def kicker(text, dark=False):
    return blurb(text, level="h4", dark=dark, title_size="14px", pad="0px", header_font=f"{KICK}|600||on|||||", header_text_color=GOLD, header_letter_spacing="2px", custom_margin="0px||14px||false|false", animation_style="none")


def title(text, body="", dark=False, h1=False, size="48px"):
    return blurb(text, body, level="h1" if h1 else "h2", dark=dark, title_size=size, body_size="19px", pad="0px", animation_style="fade")


def button(text, url, kind="gold", align="left"):
    if kind == "gold":
        tc, bg, bc = WHITE, GOLD, GOLD
    elif kind == "light":
        tc, bg, bc = BLUE, WHITE, WHITE
    else:
        tc, bg, bc = GOLD, "RGBA(255,255,255,0)", GOLD
    return sc("et_pb_button", base() | {
        "button_text": text, "button_url": url, "button_alignment": align,
        "custom_button": "on", "button_text_size": "15px", "button_text_color": tc,
        "button_bg_color": bg, "button_border_width": "1px", "button_border_color": bc,
        "button_border_radius": "0px", "button_font": f"{HEAD}|600||on|||||",
        "custom_padding": "15px|29px|15px|29px|false|false", "animation_style": "none",
        "custom_margin": "18px|10px|0px|0px|false|false",
        "button_bg_color__hover_enabled": "on|hover", "button_bg_color__hover": BLUE,
        "button_text_color__hover_enabled": "on|hover", "button_text_color__hover": WHITE,
        "button_border_color__hover_enabled": "on|hover", "button_border_color__hover": BLUE,
        "custom_css_main_element_phone": "display:block!important;width:100%!important;",
    })


def image(url, alt, height="360px", radii="on|24px|24px|24px|24px", **extra):
    return sc("et_pb_image", base() | {
        "src": url, "alt": alt, "force_fullwidth": "on", "height": height,
        "height_tablet": "320px", "height_phone": "230px", "object_fit": "cover",
        "border_radii": radii, "animation_style": "fade", "animation_duration": "700ms",
    } | extra)


def counter(number, label):
    return sc("et_pb_number_counter", base() | {
        "admin_label": f"Indicador nativo · {label}", "title": label, "number": number,
        "percent_sign": "off", "title_font": f"{BODY}|300|||||||", "title_text_color": WHITE,
        "title_font_size": "15px", "title_line_height": "1.45em", "number_font": "Poppins|300|||||||",
        "number_text_color": GOLD, "number_font_size": "44px", "number_font_size_phone": "36px",
        "background_color": "rgba(3,19,54,0.68)", "custom_padding": "18px|18px|18px|18px|false|false",
        "border_width_all": "1px", "border_color_all": "rgba(255,255,255,0.16)", "border_radii": "on|10px|10px|10px|10px",
    })


def toggle(q, a, opened=False):
    return sc("et_pb_toggle", base() | {
        "title": q, "open": "on" if opened else "off", "title_font": f"{HEAD}|600|||||||",
        "title_text_color": BLUE, "title_font_size": "20px", "body_font": f"{BODY}|300|||||||",
        "body_text_color": INK, "body_font_size": "17px", "body_line_height": "1.6em",
        "background_color": WHITE, "custom_padding": "22px|24px|22px|24px|false|false",
        "custom_margin": "0px||12px||false|false", "border_width_left": "5px",
        "border_color_left": GOLD, "border_width_all": "0px", "border_radii": "on|0px|14px|14px|0px",
        "box_shadow_style": "preset4", "box_shadow_color": "rgba(11,43,107,0.09)",
    }, f"<p>{html.escape(a)}</p>")


def contact_form():
    fields = [
        ("nombre", "Nombre", "input", "off", "Nombre y apellidos", None),
        ("telefono", "Teléfono", "input", "off", "Teléfono de contacto", None),
        ("email", "Email", "email", "off", "Email", None),
        ("tipo", "Tipo de caso", "select", "off", "", "Servicios para herencias bloqueadas|Búsqueda de herederos|Tramitación de herencia|Compra de derechos hereditarios|Me han contactado"),
        ("resumen", "Resumen", "text", "on", "Explica qué ocurre, qué documentación tienes y qué necesitas resolver.", None),
    ]
    children = ""
    for fid, label, ftype, full, placeholder, options in fields:
        a = base() | {
            "field_id": fid, "field_title": label, "field_type": ftype, "fullwidth_field": full,
            "required_mark": "off", "field_background_color": WHITE, "field_text_color": INK,
            "field_font": f"{BODY}|300|||||||", "field_font_size": "16px",
            "field_border_width_all": "1px", "field_border_color_all": "rgba(11,43,107,0.18)",
            "field_border_radius": "3px", "field_placeholder": placeholder,
            "select_options": options, "minimum_height": "140px" if fid == "resumen" else None,
        }
        children += sc("et_pb_contact_field", a)
    return sc("et_pb_contact_form", base() | {
        "email": "tecuidamos@legadia.es", "success_message": "Gracias. Hemos recibido tu consulta.",
        "submit_button_text": "Enviar consulta", "captcha": "off", "custom_button": "on",
        "button_text_size": "15px", "button_text_color": WHITE, "button_bg_color": GOLD,
        "button_border_width": "1px", "button_border_color": GOLD, "button_border_radius": "0px",
        "button_font": f"{HEAD}|600||on|||||", "form_field_background_color": WHITE,
        "form_field_text_color": INK, "form_field_font": f"{BODY}|300|||||||", "form_field_font_size": "16px",
        "background_color": WHITE, "custom_padding": "34px|34px|34px|34px|false|false",
        "border_radii": "on|20px|20px|20px|20px", "box_shadow_style": "preset4",
        "box_shadow_color": "rgba(3,19,54,0.18)",
    }, children)


def build_json():
    parts = []
    hero_left = kicker(TEXT["hero_k"], True) + title(TEXT["hero_h"], TEXT["hero_p"], True, True, "62px")
    hero_left += button("Escribir WhatsApp", "https://wa.me/34919359472?text=Hola%20Legadia%2C%20necesito%20ayuda%20con%20una%20herencia.")
    hero_left += button("Llamar ahora", "tel:+34919359472", "outline")
    hero_right = image(IMG["hero"], TEXT["hero_card"], "350px", "on|28px|28px|0px|0px")
    hero_right += blurb(TEXT["hero_label"], TEXT["hero_card"], dark=True, bg="rgba(3,19,54,0.90)", pad="22px|24px|24px|24px|false|false", title_size="14px", body_size="22px", header_text_color=GOLD, border_radii="on|0px|0px|28px|28px", animation_style="none")
    parts.append(section(row(col(hero_left, "3_5", custom_padding="0px|32px|0px|0px|false|false") + col(hero_right, "2_5", background_color="rgba(255,255,255,0.08)", custom_padding="14px|14px|14px|14px|false|false", border_width_all="1px", border_color_all="rgba(255,255,255,0.28)", border_radii="on|30px|30px|30px|30px"), "3_5,2_5", make_equal="on", use_custom_gutter="on", gutter_width="2"), "01 · Hero premium servicios", background_color=PURPLE, use_background_color_gradient="on", background_color_gradient_type="circular", background_color_gradient_direction_radial="left", background_color_gradient_stops="rgba(8,0,124,0.90) 0%|rgba(3,19,54,0.96) 64%|rgba(224,153,0,0.58) 100%", background_image=HOME_IMG, background_color_gradient_overlays_image="on", parallax="on", parallax_method="off", custom_padding="132px||104px||true|false", custom_padding_tablet="98px||80px||true|false", custom_padding_phone="70px||62px||true|false"))

    intro = row(col(kicker(TEXT["services_k"], True) + title(TEXT["services_h"], "", True, False, "52px"), "2_5") + col(blurb("", TEXT["services_p"], dark=True, body_size="20px"), "3_5"), "2_5,3_5", custom_margin="0px||54px||false|false")
    s1 = row(col(image(IMG["search"], TEXT["s1_h"], "440px", "on|24px|0px|0px|24px"), "1_2") + col(blurb("01", "", dark=True, title_size="54px", header_text_color=GOLD, animation_style="none") + blurb(TEXT["s1_h"], TEXT["s1_p"], dark=True, title_size="36px") + button("Ver servicio", "busqueda-de-herederos.html"), "1_2", background_color=BLUE, custom_padding="62px|56px|58px|56px|false|false", border_radii="on|0px|24px|24px|0px"), "1_2,1_2", make_equal="on", use_custom_gutter="on", gutter_width="1", custom_margin="0px||34px||false|false")
    s2 = row(col(blurb("02", "", dark=False, title_size="54px", header_text_color=NAVY, animation_style="none") + blurb(TEXT["s2_h"], TEXT["s2_p"], dark=False, title_size="36px", header_text_color=NAVY, body_text_color=NAVY) + button("Ver servicio", "tramitacion-de-herencias.html", "light"), "1_2", background_color=GOLD, custom_padding="62px|56px|58px|56px|false|false", border_radii="on|24px|0px|0px|24px") + col(image(IMG["process"], TEXT["s2_h"], "440px", "on|0px|24px|24px|0px"), "1_2"), "1_2,1_2", make_equal="on", use_custom_gutter="on", gutter_width="1", custom_margin="0px||34px||false|false")
    s3body = kicker("03 · Valoración y salida", True) + title(TEXT["s3_h"], TEXT["s3_p"], True, False, "48px") + button("Ver servicio", "compra-de-derechos-hereditarios.html")
    s3 = row(col(s3body, "1_1", background_color="rgba(3,19,54,0.72)", custom_padding="72px|64px|72px|64px|false|false", border_width_all="1px", border_color_all="rgba(255,255,255,0.22)", border_radii="on|24px|24px|24px|24px"), "1_1", background_image=IMG["rights"], background_color="rgba(3,19,54,0.78)", background_blend="multiply", custom_padding="36px|36px|36px|36px|false|false", border_radii="on|24px|24px|24px|24px", custom_margin="0px||72px||false|false")
    access_intro = row(col(kicker(TEXT["access_k"], True) + title(TEXT["access_h"], "", True, False, "42px"), "2_5") + col(blurb("", TEXT["access_p"], dark=True, body_size="19px"), "3_5"), "2_5,3_5", custom_margin="0px||36px||false|false")
    contacted = col(blurb(TEXT["contacted_h"], TEXT["contacted_p"], dark=True, bg="rgba(3,19,54,0.72)", pad="44px|38px|24px|38px|false|false", title_size="32px", border_radii="on|22px|22px|0px|0px") + button("Comprobar mi caso", "http://185.232.41.199/~wwwlegadianew/te-hemos-contactado/"), "1_2", background_image=IMG["contacted"], background_color="rgba(3,19,54,0.76)", background_blend="multiply", custom_padding="44px|38px|40px|38px|false|false", border_radii="on|24px|24px|24px|24px", min_height="420px")
    partners = col(blurb(TEXT["partners_h"], TEXT["partners_p"], dark=False, bg="RGBA(255,255,255,0)", pad="44px|38px|24px|38px|false|false", title_size="32px", header_text_color=NAVY, body_text_color=NAVY) + button("Ver colaboraciones", "http://185.232.41.199/~wwwlegadianew/colaboramos/", "light"), "1_2", background_color=GOLD, background_image=IMG["partners"], background_blend="soft-light", custom_padding="44px|38px|40px|38px|false|false", border_radii="on|24px|24px|24px|24px", min_height="420px")
    access = row(contacted + partners, "1_2,1_2", make_equal="on", use_custom_gutter="on", gutter_width="2")
    parts.append(section(intro + s1 + s2 + s3 + access_intro + access, "02 · Portal premium de servicios", background_color=NAVY, use_background_color_gradient="on", background_color_gradient_direction="135deg", background_color_gradient_stops=f"{NAVY} 0%|{BLUE} 72%|rgba(224,153,0,0.48) 145%", custom_padding="108px||118px||true|false", custom_padding_tablet="84px||92px||true|false", custom_padding_phone="66px||74px||true|false"))

    d_intro = row(col(kicker(TEXT["diag_k"]) + title(TEXT["diag_h"], "", False, False, "48px"), "2_5") + col(blurb("", TEXT["diag_p"], body_size="20px"), "3_5"), "2_5,3_5", custom_margin="0px||48px||false|false")
    dcols = ""
    for i, value in enumerate(TEXT["diag"], 1):
        selected = i == 3
        dcols += col(blurb(f"0{i}", value, dark=selected, bg=GOLD if selected else "RGBA(255,255,255,0)", pad="34px|26px|34px|26px|false|false", title_size="38px", body_size="18px", header_text_color=NAVY if selected else GOLD, body_text_color=NAVY if selected else INK, border_width_left="1px", border_color_left=GOLD if selected else "rgba(11,43,107,0.18)", animation_delay=f"{i*80}ms"), "1_4")
    parts.append(section(d_intro + row(dcols, "1_4,1_4,1_4,1_4", make_equal="on", use_custom_gutter="on", gutter_width="1"), "03 · Diagnóstico editorial sin cajas", background_color=SOFT, use_background_color_gradient="on", background_color_gradient_direction="135deg", background_color_gradient_stops=f"{SOFT} 0%|#FFFFFF 72%|rgba(224,153,0,0.14) 100%", custom_padding="94px||98px||true|false"))

    g_intro = row(col(kicker(TEXT["gallery_k"]) + title(TEXT["gallery_h"], "", False, False, "48px"), "3_5") + col(blurb("", TEXT["gallery_p"], body_size="19px"), "2_5"), "3_5,2_5", custom_margin="0px||42px||false|false")
    big = col(image(IMG["team"], TEXT["gallery"][0], "470px", "on|24px|24px|0px|0px") + blurb("01", TEXT["gallery"][0], dark=True, bg=NAVY, pad="22px|24px|24px|24px|false|false", title_size="20px", header_text_color=GOLD, border_radii="on|0px|0px|24px|24px"), "1_2")
    small1 = col(image(IMG["docs"], TEXT["gallery"][1], "300px", "on|20px|20px|0px|0px") + blurb("02", TEXT["gallery"][1], dark=False, bg=GOLD, pad="20px|20px|22px|20px|false|false", title_size="18px", header_text_color=NAVY, body_text_color=NAVY, border_radii="on|0px|0px|20px|20px"), "1_4")
    small2 = col(image(IMG["talk"], TEXT["gallery"][2], "300px", "on|20px|20px|0px|0px") + blurb("03", TEXT["gallery"][2], dark=True, bg=BLUE, pad="20px|20px|22px|20px|false|false", title_size="18px", header_text_color=GOLD, border_radii="on|0px|0px|20px|20px"), "1_4")
    parts.append(section(g_intro + row(big + small1 + small2, "1_2,1_4,1_4", make_equal="on", use_custom_gutter="on", gutter_width="2"), "04 · Galería asimétrica premium", background_color=WHITE, custom_padding="94px||98px||true|false"))

    class_left = image(IMG["hero"], TEXT["class_h"], "100%", "on|28px|0px|0px|28px", min_height="560px")
    class_right = kicker(TEXT["class_k"], True) + title(TEXT["class_h"], TEXT["class_p1"], True, False, "48px") + blurb("", TEXT["class_p2"], dark=True, body_size="18px", custom_margin="0px||18px||false|false") + blurb(TEXT["class_q"], "", dark=False, bg=GOLD, pad="22px|24px|22px|24px|false|false", title_size="20px", header_text_color=NAVY, border_radii="on|14px|14px|14px|14px", animation_style="fade")
    parts.append(section(row(col(class_left, "1_2", overflow_x="hidden", overflow_y="hidden") + col(class_right, "1_2", background_color=BLUE, custom_padding="62px|56px|62px|56px|false|false", border_radii="on|0px|28px|28px|0px"), "1_2,1_2", make_equal="on", use_custom_gutter="on", gutter_width="1"), "05 · Clasificación editorial", background_color=WHITE, custom_padding="96px||96px||true|false"))

    m_intro = row(col(kicker(TEXT["method_k"], True) + title(TEXT["method_h"], "", True, False, "48px"), "2_5") + col(blurb("", TEXT["method_p"], dark=True, body_size="20px"), "3_5"), "2_5,3_5", custom_margin="0px||58px||false|false")
    mcols = ""
    for i, (h, p) in enumerate(TEXT["method"]):
        mcols += col(blurb(h, p, dark=True, bg="RGBA(255,255,255,0)", pad="26px|22px|28px|22px|false|false", title_size="22px", border_width_top="3px", border_color_top=GOLD if i == 2 else "rgba(255,255,255,0.28)", animation_delay=f"{i*100}ms"), "1_4")
    parts.append(section(m_intro + row(mcols, "1_4,1_4,1_4,1_4", make_equal="on", use_custom_gutter="on", gutter_width="2"), "06 · Metodología lineal premium", background_color=NAVY, background_image=IMG["docs"], background_blend="multiply", use_background_color_gradient="on", background_color_gradient_direction="90deg", background_color_gradient_stops="rgba(3,19,54,0.96) 0%|rgba(11,43,107,0.88) 100%", background_color_gradient_overlays_image="on", parallax="on", parallax_method="off", custom_padding="104px||108px||true|false"))

    aud = kicker(TEXT["aud_k"]) + title(TEXT["aud_h"], TEXT["aud_p"], False, False, "46px")
    for i, value in enumerate(TEXT["aud"], 1):
        aud += blurb(f"0{i}", value, icon="N", bg="RGBA(255,255,255,0)", pad="14px|0px|14px|0px|false|false", title_size="18px", body_size="17px", icon_placement="left", use_icon_font_size="on", icon_font_size="20px", custom_margin="0px||0px||false|false", border_width_bottom="1px", border_color_bottom="rgba(11,43,107,0.12)")
    err = kicker(TEXT["err_k"], True) + title(TEXT["err_h"], TEXT["err_p"], True, False, "46px") + blurb(TEXT["err_sub"], "", dark=True, title_size="23px", custom_margin="16px||12px||false|false")
    for value in TEXT["err"]:
        err += blurb("×", value, dark=True, bg="rgba(3,19,54,0.58)", pad="14px|16px|14px|16px|false|false", title_size="22px", body_size="16px", header_text_color=GOLD, custom_margin="0px||8px||false|false", border_width_left="3px", border_color_left=GOLD)
    parts.append(section(row(col(aud, "1_2", custom_padding="0px|34px|0px|0px|false|false") + col(err, "1_2", background_color="rgba(3,19,54,0.84)", background_image=IMG["scene"], background_blend="multiply", custom_padding="54px|46px|54px|46px|false|false", border_radii="on|26px|26px|26px|26px"), "1_2,1_2", make_equal="on", use_custom_gutter="on", gutter_width="2"), "07 · Público y errores contrastados", background_color=WHITE, custom_padding="98px||102px||true|false"))

    scen_intro = row(col(kicker(TEXT["scen_k"], True) + title(TEXT["scen_h"], "", True, False, "48px"), "2_5") + col(blurb("", TEXT["scen_p"], dark=True, body_size="19px"), "3_5"), "2_5,3_5", custom_margin="0px||42px||false|false")
    scen_cols = ""
    for i, value in enumerate(TEXT["scen"], 1):
        scen_cols += col(blurb(f"0{i}", value, dark=True, bg="rgba(255,255,255,0.06)", pad="26px|22px|26px|22px|false|false", title_size="32px", header_text_color=GOLD, border_width_all="1px", border_color_all="rgba(255,255,255,0.18)", border_radii="on|16px|16px|16px|16px"), "1_4")
    doc_intro = row(col(kicker(TEXT["doc_k"], True) + title(TEXT["doc_h"], "", True, False, "44px"), "3_5") + col(blurb("", TEXT["doc_p"], dark=True, body_size="18px"), "2_5"), "3_5,2_5", custom_margin="72px||34px||false|false")
    doc_cols = ""
    for i, value in enumerate(TEXT["docs"], 1):
        doc_cols += col(blurb(str(i), value, dark=True, bg="rgba(3,19,54,0.48)", pad="20px|18px|20px|18px|false|false", title_size="26px", header_text_color=GOLD, border_width_top="1px", border_color_top="rgba(255,255,255,0.24)"), "1_4")
    parts.append(section(scen_intro + row(scen_cols, "1_4,1_4,1_4,1_4", make_equal="on", use_custom_gutter="on", gutter_width="2") + doc_intro + row(doc_cols, "1_4,1_4,1_4,1_4", make_equal="on", use_custom_gutter="on", gutter_width="2"), "08 · Escenarios y documentación parallax", background_color=BLUE, background_image=IMG["scene"], background_blend="multiply", use_background_color_gradient="on", background_color_gradient_direction="135deg", background_color_gradient_stops="rgba(11,43,107,0.92) 0%|rgba(36,52,103,0.92) 68%|rgba(224,153,0,0.44) 140%", background_color_gradient_overlays_image="on", parallax="on", parallax_method="off", custom_padding="104px||110px||true|false"))

    crit = kicker(TEXT["crit_k"]) + title(TEXT["crit_h"], TEXT["crit_p1"], False, False, "46px") + blurb("", TEXT["crit_p2"], body_size="18px")
    trust = kicker(TEXT["trust_k"], True) + title(TEXT["trust_h"], TEXT["trust_p"], True, False, "46px") + blurb(TEXT["trust_card_h"], TEXT["trust_card_p"], dark=True, bg="rgba(3,19,54,0.58)", pad="22px|22px|22px|22px|false|false", title_size="23px", border_radii="on|14px|14px|14px|14px")
    trust += counter("+400", "Herederos encontrados en expedientes complejos.") + counter("3", "Áreas conectadas: análisis jurídico, documentación e investigación genealógica.") + counter("1", "Interlocutor claro para ordenar el siguiente paso del expediente.")
    parts.append(section(row(col(crit, "1_2", custom_padding="54px|44px|54px|44px|false|false") + col(trust, "1_2", background_image=IMG["team"], background_color="rgba(3,19,54,0.88)", background_blend="multiply", custom_padding="54px|44px|54px|44px|false|false", border_radii="on|26px|26px|26px|26px"), "1_2,1_2", make_equal="on", use_custom_gutter="on", gutter_width="2"), "09 · Objetivo y confianza premium", background_color=SOFT, custom_padding="98px||102px||true|false"))

    direct = kicker(TEXT["form_k"], True) + title(TEXT["form_h"], TEXT["form_p"], True, False, "44px") + blurb(TEXT["direct_h"], TEXT["direct_p"], dark=True, title_size="25px", custom_margin="18px||20px||false|false")
    direct += blurb("+34 91 935 94 72", "Llamar ahora", dark=True, bg="rgba(255,255,255,0.08)", pad="16px|18px|16px|18px|false|false", title_size="20px", border_width_all="1px", border_color_all="rgba(255,255,255,0.18)", url="tel:+34919359472")
    direct += blurb("tecuidamos@legadia.es", "Escribir por email", dark=True, bg="rgba(255,255,255,0.08)", pad="16px|18px|16px|18px|false|false", title_size="20px", border_width_all="1px", border_color_all="rgba(255,255,255,0.18)", url="mailto:tecuidamos@legadia.es")
    direct += button("Escribir por WhatsApp", "https://wa.me/34919359472?text=Hola%20Legadia%2C%20necesito%20ayuda%20con%20una%20herencia.")
    parts.append(section(row(col(direct, "2_5", background_color=NAVY, custom_padding="54px|44px|54px|44px|false|false", border_radii="on|24px|0px|0px|24px") + col(contact_form(), "3_5", background_color=GOLD, custom_padding="36px|36px|36px|36px|false|false", border_radii="on|0px|24px|24px|0px"), "2_5,3_5", make_equal="on", use_custom_gutter="on", gutter_width="1"), "10 · Consulta inicial nativa", background_color=GOLD, use_background_color_gradient="on", background_color_gradient_direction="135deg", background_color_gradient_stops=f"{GOLD} 0%|#F4B93E 100%", custom_padding="94px||98px||true|false"))

    faq_left = kicker(TEXT["faq_k"]) + title(TEXT["faq_h"], TEXT["faq_p"], False, False, "46px")
    faq_right = "".join(toggle(q, a, i == 0) for i, (q, a) in enumerate(TEXT["faq"]))
    parts.append(section(row(col(faq_left, "2_5", custom_padding="0px|34px|0px|0px|false|false") + col(faq_right, "3_5"), "2_5,3_5", use_custom_gutter="on", gutter_width="2"), "11 · FAQ nativa minimalista", background_color=WHITE, custom_padding="96px||100px||true|false"))

    final = kicker(TEXT["final_k"], True) + title(TEXT["final_h"], TEXT["final_p"], True, False, "52px") + button("Contar mi caso", "#consulta", "gold", "center") + button("Llamar ahora", "tel:+34919359472", "light", "center")
    parts.append(section(row(col(final, "1_1"), "1_1", max_width="920px", text_orientation="center"), "12 · Cierre premium parallax", background_color=NAVY, background_image=IMG["talk"], background_blend="multiply", use_background_color_gradient="on", background_color_gradient_direction="90deg", background_color_gradient_stops="rgba(3,19,54,0.94) 0%|rgba(8,0,124,0.78) 58%|rgba(224,153,0,0.62) 140%", background_color_gradient_overlays_image="on", parallax="on", parallax_method="off", custom_padding="104px||108px||true|false"))
    markup = "".join(parts)
    return {"context": "et_builder", "data": {"990001": markup}}


def build_html():
    services = [
        ("01", TEXT["s1_h"], TEXT["s1_p"], IMG["search"], "busqueda-de-herederos.html", "navy"),
        ("02", TEXT["s2_h"], TEXT["s2_p"], IMG["process"], "tramitacion-de-herencias.html", "gold"),
        ("03", TEXT["s3_h"], TEXT["s3_p"], IMG["rights"], "compra-de-derechos-hereditarios.html", "image"),
    ]
    def esc(x): return html.escape(str(x), quote=True)
    service_html = ""
    for num, h, p, img, url, style in services:
        service_html += f'<article class="service {style}" style="--image:url({esc(img)})"><div class="service-image"><img src="{esc(img)}" alt="{esc(h)}"></div><div class="service-copy"><span>{num}</span><h3>{esc(h)}</h3><p>{esc(p)}</p><a href="{esc(url)}">Ver servicio</a></div></article>'
    diag = "".join(f'<article class="diag {"featured" if i==2 else ""}"><b>0{i+1}</b><p>{esc(x)}</p></article>' for i, x in enumerate(TEXT["diag"]))
    methods = "".join(f'<article class="method {"featured" if i==2 else ""}"><b>{esc(h)}</b><p>{esc(p)}</p></article>' for i,(h,p) in enumerate(TEXT["method"]))
    audience = "".join(f'<li><b>0{i+1}</b>{esc(x)}</li>' for i,x in enumerate(TEXT["aud"]))
    errors = "".join(f'<li><b>×</b>{esc(x)}</li>' for x in TEXT["err"])
    scenarios = "".join(f'<article><b>0{i+1}</b><p>{esc(x)}</p></article>' for i,x in enumerate(TEXT["scen"]))
    docs = "".join(f'<article><b>{i+1}</b><p>{esc(x)}</p></article>' for i,x in enumerate(TEXT["docs"]))
    faqs = "".join(f'<details {"open" if i==0 else ""}><summary>{esc(q)}</summary><p>{esc(a)}</p></details>' for i,(q,a) in enumerate(TEXT["faq"]))
    options = "".join(f'<option>{esc(x)}</option>' for x in ["Servicios para herencias bloqueadas","Búsqueda de herederos","Tramitación de herencia","Compra de derechos hereditarios","Me han contactado"])
    css = f'''@import url('https://fonts.googleapis.com/css2?family=Catamaran:wght@600&family=Josefin+Sans:wght@300;400;600&family=Poppins:wght@300&display=swap');
:root{{--navy:{NAVY};--blue:{BLUE};--purple:{PURPLE};--violet:{VIOLET};--gold:{GOLD};--white:#fff;--soft:{SOFT};--ink:{INK};--muted:{MUTED}}}*{{box-sizing:border-box}}html{{scroll-behavior:smooth}}body{{margin:0;font-family:'Josefin Sans',sans-serif;color:var(--muted);line-height:1.58;background:#fff}}img{{display:block;width:100%;object-fit:cover}}a{{text-decoration:none}}.wrap{{width:min(1240px,calc(100% - 42px));margin:auto}}.k{{font:600 14px Catamaran,sans-serif;letter-spacing:2px;text-transform:uppercase;color:var(--gold);margin-bottom:14px}}h1,h2,h3{{margin:0;font-weight:300;line-height:1.1}}h1{{font-size:clamp(40px,5.2vw,68px);color:#fff}}h2{{font-size:clamp(34px,4vw,52px);color:var(--blue)}}h3{{font-size:clamp(27px,3vw,38px)}}p{{font-size:18px;margin:16px 0 0}}.dark h2,.dark h3,.dark p{{color:#fff}}.btn{{display:inline-flex;padding:15px 29px;border:1px solid var(--gold);background:var(--gold);color:#fff;font-weight:600;margin:24px 10px 0 0}}.btn.light{{background:#fff;color:var(--blue);border-color:#fff}}.hero{{padding:132px 0 104px;background:radial-gradient(circle at left,rgba(8,0,124,.9),rgba(3,19,54,.96) 64%,rgba(224,153,0,.58)),url('{HOME_IMG}') center/cover fixed}}.hero-grid{{display:grid;grid-template-columns:3fr 2fr;gap:42px;align-items:center}}.hero p{{color:rgba(255,255,255,.82);font-size:20px}}.hero-visual{{padding:14px;border:1px solid rgba(255,255,255,.28);border-radius:30px;background:rgba(255,255,255,.08)}}.hero-visual img{{height:350px;border-radius:28px 28px 0 0}}.hero-note{{background:rgba(3,19,54,.9);padding:22px 24px 24px;border-radius:0 0 28px 28px}}.hero-note small{{color:var(--gold);font:600 13px Catamaran}}.hero-note h3{{font-size:24px;color:#fff;margin-top:8px}}section{{padding:96px 0}}.section-head{{display:grid;grid-template-columns:2fr 3fr;gap:42px;align-items:end;margin-bottom:52px}}.section-head p{{font-size:20px}}.services{{background:linear-gradient(135deg,var(--navy),var(--blue) 72%,rgba(224,153,0,.48) 145%);padding:108px 0 118px}}.service{{display:grid;grid-template-columns:1fr 1fr;margin-bottom:34px;border-radius:24px;overflow:hidden}}.service-image img{{height:440px}}.service-copy{{padding:62px 56px;display:flex;flex-direction:column;justify-content:center}}.service-copy span{{font:300 54px Poppins;color:var(--gold)}}.service-copy p{{font-size:18px}}.service-copy a{{display:inline-flex;align-self:flex-start;margin-top:24px;padding:15px 29px;background:var(--gold);color:#fff;font-weight:600}}.service.navy .service-copy{{background:var(--blue);color:#fff}}.service.navy h3,.service.navy p{{color:#fff}}.service.gold{{grid-template-areas:'copy image'}}.service.gold .service-copy{{grid-area:copy;background:var(--gold);color:var(--navy)}}.service.gold .service-image{{grid-area:image}}.service.gold .service-copy span{{color:var(--navy)}}.service.gold .service-copy a{{background:#fff;color:var(--blue)}}.service.image{{grid-template-columns:1fr;background:linear-gradient(rgba(3,19,54,.75),rgba(3,19,54,.75)),var(--image) center/cover}}.service.image .service-image{{display:none}}.service.image .service-copy{{max-width:650px;margin:36px;padding:72px 64px;background:rgba(3,19,54,.72);border:1px solid rgba(255,255,255,.22);border-radius:24px;color:#fff}}.service.image h3,.service.image p{{color:#fff}}.access-head{{margin-top:72px}}.access-grid{{display:grid;grid-template-columns:1fr 1fr;gap:20px}}.access{{min-height:420px;padding:44px 38px;border-radius:24px;display:flex;flex-direction:column;justify-content:flex-end}}.access.contact{{background:linear-gradient(rgba(3,19,54,.76),rgba(3,19,54,.76)),url('{IMG['contacted']}') center/cover;color:#fff}}.access.partner{{background:linear-gradient(rgba(224,153,0,.82),rgba(224,153,0,.82)),url('{IMG['partners']}') center/cover;color:var(--navy)}}.access h3{{font-size:32px}}.access p{{font-size:17px}}.diagnosis{{background:linear-gradient(135deg,var(--soft),#fff 72%,rgba(224,153,0,.14))}}.diag-grid{{display:grid;grid-template-columns:repeat(4,1fr)}}.diag{{padding:34px 26px;border-left:1px solid rgba(11,43,107,.18)}}.diag b{{font:300 38px Poppins;color:var(--gold)}}.diag.featured{{background:var(--gold);color:var(--navy)}}.diag.featured b{{color:var(--navy)}}.gallery-grid{{display:grid;grid-template-columns:2fr 1fr 1fr;gap:20px}}.visual img{{height:300px;border-radius:20px 20px 0 0}}.visual.big img{{height:470px;border-radius:24px 24px 0 0}}.visual figcaption{{margin:0;background:var(--blue);color:#fff;padding:20px}}.visual.big figcaption{{background:var(--navy)}}.visual.gold figcaption{{background:var(--gold);color:var(--navy)}}.classify-grid{{display:grid;grid-template-columns:1fr 1fr}}.classify-grid>img{{height:100%;min-height:560px;border-radius:28px 0 0 28px}}.classify-copy{{background:var(--blue);padding:62px 56px;border-radius:0 28px 28px 0;color:#fff}}.quote{{background:var(--gold);color:var(--navy);padding:22px 24px;border-radius:14px;font-size:20px;margin-top:24px}}.methodology{{background:linear-gradient(90deg,rgba(3,19,54,.96),rgba(11,43,107,.88)),url('{IMG['docs']}') center/cover fixed}}.method-grid{{display:grid;grid-template-columns:repeat(4,1fr);gap:20px}}.method{{padding:26px 22px;border-top:3px solid rgba(255,255,255,.28);color:#fff}}.method.featured{{border-color:var(--gold)}}.method b{{font-size:22px;color:#fff}}.audience-grid{{display:grid;grid-template-columns:1fr 1fr;gap:30px}}.editorial-list{{list-style:none;padding:0;margin:28px 0 0}}.editorial-list li{{padding:16px 0;border-bottom:1px solid rgba(11,43,107,.12);font-size:17px}}.editorial-list b{{color:var(--gold);margin-right:14px}}.errors{{padding:54px 46px;border-radius:26px;background:linear-gradient(rgba(3,19,54,.84),rgba(3,19,54,.84)),url('{IMG['scene']}') center/cover;color:#fff}}.errors ul{{list-style:none;padding:0}}.errors li{{display:flex;gap:14px;padding:14px 16px;margin-top:8px;background:rgba(3,19,54,.58);border-left:3px solid var(--gold)}}.errors b{{color:var(--gold);font-size:22px}}.scenarios{{background:linear-gradient(135deg,rgba(11,43,107,.92),rgba(36,52,103,.92) 68%,rgba(224,153,0,.44) 140%),url('{IMG['scene']}') center/cover fixed}}.scenario-grid,.doc-grid{{display:grid;grid-template-columns:repeat(4,1fr);gap:18px}}.scenario-grid article{{padding:26px 22px;background:rgba(255,255,255,.06);border:1px solid rgba(255,255,255,.18);border-radius:16px;color:#fff}}.scenario-grid b,.doc-grid b{{font:300 32px Poppins;color:var(--gold)}}.doc-head{{margin-top:72px}}.doc-grid article{{padding:20px 18px;border-top:1px solid rgba(255,255,255,.24);color:#fff}}.trust-grid{{display:grid;grid-template-columns:1fr 1fr;gap:20px}}.criterion{{padding:54px 44px}}.trust-panel{{padding:54px 44px;border-radius:26px;background:linear-gradient(rgba(3,19,54,.88),rgba(3,19,54,.88)),url('{IMG['team']}') center/cover;color:#fff}}.counters{{display:grid;grid-template-columns:repeat(3,1fr);gap:10px;margin-top:22px}}.counter{{padding:18px;background:rgba(3,19,54,.68);border:1px solid rgba(255,255,255,.16);border-radius:10px}}.counter b{{font:300 44px Poppins;color:var(--gold)}}.counter span{{display:block;font-size:15px}}.contact-section{{background:linear-gradient(135deg,var(--gold),#F4B93E)}}.contact-grid{{display:grid;grid-template-columns:2fr 3fr}}.direct{{padding:54px 44px;background:var(--navy);border-radius:24px 0 0 24px;color:#fff}}.direct a.channel{{display:block;color:#fff;border:1px solid rgba(255,255,255,.18);padding:16px 18px;margin-top:12px}}.form{{padding:36px;background:var(--gold);border-radius:0 24px 24px 0}}form{{background:#fff;padding:34px;border-radius:20px;box-shadow:0 22px 64px rgba(3,19,54,.18)}}.fields{{display:grid;grid-template-columns:1fr 1fr;gap:14px}}label{{display:grid;gap:7px;color:var(--blue);font-weight:600}}label.full{{grid-column:1/-1}}input,select,textarea{{padding:15px;border:1px solid rgba(11,43,107,.18);font:16px 'Josefin Sans'}}textarea{{min-height:140px}}button{{padding:15px 29px;background:var(--gold);color:#fff;border:0;margin-top:18px;font-weight:600}}.faq-grid{{display:grid;grid-template-columns:2fr 3fr;gap:34px}}details{{background:#fff;border-left:5px solid var(--gold);box-shadow:0 12px 34px rgba(11,43,107,.09);margin-bottom:12px;border-radius:0 14px 14px 0;padding:0 24px}}summary{{padding:22px 0;color:var(--blue);font-size:20px;font-weight:600;cursor:pointer}}details p{{padding-bottom:22px}}.final{{text-align:center;background:linear-gradient(90deg,rgba(3,19,54,.94),rgba(8,0,124,.78) 58%,rgba(224,153,0,.62) 140%),url('{IMG['talk']}') center/cover fixed}}.final .wrap{{max-width:920px}}@media(max-width:980px){{.hero-grid,.section-head,.service,.access-grid,.classify-grid,.audience-grid,.trust-grid,.contact-grid,.faq-grid{{grid-template-columns:1fr}}.service.gold{{grid-template-areas:'image' 'copy'}}.service.image .service-copy{{margin:20px}}.diag-grid,.method-grid,.scenario-grid,.doc-grid{{grid-template-columns:1fr 1fr}}.gallery-grid{{grid-template-columns:1fr 1fr}}.visual.big{{grid-column:1/-1}}.classify-grid>img{{border-radius:28px 28px 0 0}}.classify-copy{{border-radius:0 0 28px 28px}}.contact-grid>div{{border-radius:24px}}}}@media(max-width:680px){{.wrap{{width:calc(100% - 28px)}}section{{padding:66px 0}}.hero{{padding:72px 0 62px}}.hero-visual img{{height:250px}}.diag-grid,.method-grid,.scenario-grid,.doc-grid,.gallery-grid,.counters,.fields{{grid-template-columns:1fr}}.visual.big{{grid-column:auto}}.service-copy,.service.image .service-copy{{padding:38px 26px;margin:0}}.service-image img{{height:260px}}.access{{min-height:360px}}.btn{{width:100%;justify-content:center}}label.full{{grid-column:auto}}}}'''
    body = f'''<main>
<section class="hero dark"><div class="wrap hero-grid"><div><div class="k">{esc(TEXT['hero_k'])}</div><h1>{esc(TEXT['hero_h'])}</h1><p>{esc(TEXT['hero_p'])}</p><a class="btn" href="https://wa.me/34919359472?text=Hola%20Legadia%2C%20necesito%20ayuda%20con%20una%20herencia.">Escribir WhatsApp</a><a class="btn light" href="tel:+34919359472">Llamar ahora</a></div><div class="hero-visual"><img src="{esc(IMG['hero'])}" alt="{esc(TEXT['hero_card'])}"><div class="hero-note"><small>{esc(TEXT['hero_label'])}</small><h3>{esc(TEXT['hero_card'])}</h3></div></div></div></section>
<section class="services dark"><div class="wrap"><div class="section-head"><div><div class="k">{esc(TEXT['services_k'])}</div><h2>{esc(TEXT['services_h'])}</h2></div><p>{esc(TEXT['services_p'])}</p></div>{service_html}<div class="section-head access-head"><div><div class="k">{esc(TEXT['access_k'])}</div><h2>{esc(TEXT['access_h'])}</h2></div><p>{esc(TEXT['access_p'])}</p></div><div class="access-grid"><article class="access contact"><h3>{esc(TEXT['contacted_h'])}</h3><p>{esc(TEXT['contacted_p'])}</p><a class="btn" href="http://185.232.41.199/~wwwlegadianew/te-hemos-contactado/">Comprobar mi caso</a></article><article class="access partner"><h3>{esc(TEXT['partners_h'])}</h3><p>{esc(TEXT['partners_p'])}</p><a class="btn light" href="http://185.232.41.199/~wwwlegadianew/colaboramos/">Ver colaboraciones</a></article></div></div></section>
<section class="diagnosis"><div class="wrap"><div class="section-head"><div><div class="k">{esc(TEXT['diag_k'])}</div><h2>{esc(TEXT['diag_h'])}</h2></div><p>{esc(TEXT['diag_p'])}</p></div><div class="diag-grid">{diag}</div></div></section>
<section><div class="wrap"><div class="section-head"><div><div class="k">{esc(TEXT['gallery_k'])}</div><h2>{esc(TEXT['gallery_h'])}</h2></div><p>{esc(TEXT['gallery_p'])}</p></div><div class="gallery-grid"><figure class="visual big"><img src="{esc(IMG['team'])}" alt="{esc(TEXT['gallery'][0])}"><figcaption>01 · {esc(TEXT['gallery'][0])}</figcaption></figure><figure class="visual gold"><img src="{esc(IMG['docs'])}" alt="{esc(TEXT['gallery'][1])}"><figcaption>02 · {esc(TEXT['gallery'][1])}</figcaption></figure><figure class="visual"><img src="{esc(IMG['talk'])}" alt="{esc(TEXT['gallery'][2])}"><figcaption>03 · {esc(TEXT['gallery'][2])}</figcaption></figure></div></div></section>
<section><div class="wrap classify-grid"><img src="{esc(IMG['hero'])}" alt="{esc(TEXT['class_h'])}"><div class="classify-copy dark"><div class="k">{esc(TEXT['class_k'])}</div><h2>{esc(TEXT['class_h'])}</h2><p>{esc(TEXT['class_p1'])}</p><p>{esc(TEXT['class_p2'])}</p><div class="quote">{esc(TEXT['class_q'])}</div></div></div></section>
<section class="methodology dark"><div class="wrap"><div class="section-head"><div><div class="k">{esc(TEXT['method_k'])}</div><h2>{esc(TEXT['method_h'])}</h2></div><p>{esc(TEXT['method_p'])}</p></div><div class="method-grid">{methods}</div></div></section>
<section><div class="wrap audience-grid"><div><div class="k">{esc(TEXT['aud_k'])}</div><h2>{esc(TEXT['aud_h'])}</h2><p>{esc(TEXT['aud_p'])}</p><ul class="editorial-list">{audience}</ul></div><div class="errors dark"><div class="k">{esc(TEXT['err_k'])}</div><h2>{esc(TEXT['err_h'])}</h2><p>{esc(TEXT['err_p'])}</p><h3>{esc(TEXT['err_sub'])}</h3><ul>{errors}</ul></div></div></section>
<section class="scenarios dark"><div class="wrap"><div class="section-head"><div><div class="k">{esc(TEXT['scen_k'])}</div><h2>{esc(TEXT['scen_h'])}</h2></div><p>{esc(TEXT['scen_p'])}</p></div><div class="scenario-grid">{scenarios}</div><div class="section-head doc-head"><div><div class="k">{esc(TEXT['doc_k'])}</div><h2>{esc(TEXT['doc_h'])}</h2></div><p>{esc(TEXT['doc_p'])}</p></div><div class="doc-grid">{docs}</div></div></section>
<section style="background:{SOFT}"><div class="wrap trust-grid"><div class="criterion"><div class="k">{esc(TEXT['crit_k'])}</div><h2>{esc(TEXT['crit_h'])}</h2><p>{esc(TEXT['crit_p1'])}</p><p>{esc(TEXT['crit_p2'])}</p></div><div class="trust-panel dark"><div class="k">{esc(TEXT['trust_k'])}</div><h2>{esc(TEXT['trust_h'])}</h2><p>{esc(TEXT['trust_p'])}</p><h3>{esc(TEXT['trust_card_h'])}</h3><p>{esc(TEXT['trust_card_p'])}</p><div class="counters"><div class="counter"><b>+400</b><span>Herederos encontrados en expedientes complejos.</span></div><div class="counter"><b>3</b><span>Áreas conectadas: análisis jurídico, documentación e investigación genealógica.</span></div><div class="counter"><b>1</b><span>Interlocutor claro para ordenar el siguiente paso del expediente.</span></div></div></div></div></section>
<section id="consulta" class="contact-section"><div class="wrap contact-grid"><div class="direct dark"><div class="k">{esc(TEXT['form_k'])}</div><h2>{esc(TEXT['form_h'])}</h2><p>{esc(TEXT['form_p'])}</p><h3>{esc(TEXT['direct_h'])}</h3><p>{esc(TEXT['direct_p'])}</p><a class="channel" href="tel:+34919359472">+34 91 935 94 72 · Llamar ahora</a><a class="channel" href="mailto:tecuidamos@legadia.es">tecuidamos@legadia.es · Escribir por email</a><a class="btn" href="https://wa.me/34919359472">Escribir por WhatsApp</a></div><div class="form"><form><div class="fields"><label>Nombre<input placeholder="Nombre y apellidos"></label><label>Teléfono<input placeholder="Teléfono de contacto"></label><label>Email<input placeholder="Email"></label><label>Tipo de caso<select>{options}</select></label><label class="full">Resumen<textarea placeholder="Explica qué ocurre, qué documentación tienes y qué necesitas resolver."></textarea></label></div><button type="button">Enviar consulta</button></form></div></div></section>
<section><div class="wrap faq-grid"><div><div class="k">{esc(TEXT['faq_k'])}</div><h2>{esc(TEXT['faq_h'])}</h2><p>{esc(TEXT['faq_p'])}</p></div><div>{faqs}</div></div></section>
<section class="final dark"><div class="wrap"><div class="k">{esc(TEXT['final_k'])}</div><h2>{esc(TEXT['final_h'])}</h2><p>{esc(TEXT['final_p'])}</p><a class="btn" href="#consulta">Contar mi caso</a><a class="btn light" href="tel:+34919359472">Llamar ahora</a></div></section>
</main>'''
    return f'<!doctype html><html lang="es"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Servicios especializados para desbloquear herencias | Legadia</title><style>{css}</style></head><body>{body}</body></html>'


def build_docx(path):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LEGADIA · SERVICIOS")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(11, 43, 107)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Textos exactos utilizados en el diseño premium Divi 4")
    r.italic = True
    r.font.color.rgb = RGBColor(224, 153, 0)
    ordered = [
        (TEXT["hero_k"], TEXT["hero_h"], [TEXT["hero_p"], TEXT["hero_label"], TEXT["hero_card"]),
        (TEXT["services_k"], TEXT["services_h"], [TEXT["services_p"], TEXT["s1_h"], TEXT["s1_p"], TEXT["s2_h"], TEXT["s2_p"], TEXT["s3_h"], TEXT["s3_p"]),
        (TEXT["access_k"], TEXT["access_h"], [TEXT["access_p"], TEXT["contacted_h"], TEXT["contacted_p"], TEXT["partners_h"], TEXT["partners_p"]),
        (TEXT["diag_k"], TEXT["diag_h"], [TEXT["diag_p"]] + TEXT["diag"]),
        (TEXT["gallery_k"], TEXT["gallery_h"], [TEXT["gallery_p"]] + TEXT["gallery"]),
        (TEXT["class_k"], TEXT["class_h"], [TEXT["class_p1"], TEXT["class_p2"], TEXT["class_q"]),
        (TEXT["method_k"], TEXT["method_h"], [TEXT["method_p"]] + [f"{a}: {b}" for a,b in TEXT["method"]]),
        (TEXT["aud_k"], TEXT["aud_h"], [TEXT["aud_p"]] + TEXT["aud"]),
        (TEXT["err_k"], TEXT["err_h"], [TEXT["err_p"], TEXT["err_sub"]] + TEXT["err"]),
        (TEXT["scen_k"], TEXT["scen_h"], [TEXT["scen_p"]] + TEXT["scen"]),
        (TEXT["doc_k"], TEXT["doc_h"], [TEXT["doc_p"]] + TEXT["docs"]),
        (TEXT["crit_k"], TEXT["crit_h"], [TEXT["crit_p1"], TEXT["crit_p2"]]),
        (TEXT["trust_k"], TEXT["trust_h"], [TEXT["trust_p"], TEXT["trust_card_h"], TEXT["trust_card_p"], "+400 Herederos encontrados en expedientes complejos.", "3 Áreas conectadas: análisis jurídico, documentación e investigación genealógica.", "1 Interlocutor claro para ordenar el siguiente paso del expediente."]),
        (TEXT["form_k"], TEXT["form_h"], [TEXT["form_p"], TEXT["direct_h"], TEXT["direct_p"], "+34 91 935 94 72", "tecuidamos@legadia.es", "Nombre y apellidos", "Teléfono de contacto", "Email", "Tipo de caso", "Servicios para herencias bloqueadas", "Búsqueda de herederos", "Tramitación de herencia", "Compra de derechos hereditarios", "Me han contactado", "Explica qué ocurre, qué documentación tienes y qué necesitas resolver.", "Enviar consulta"]),
        (TEXT["faq_k"], TEXT["faq_h"], [TEXT["faq_p"]] + [f"{q}: {a}" for q,a in TEXT["faq"]]),
        (TEXT["final_k"], TEXT["final_h"], [TEXT["final_p"], "Contar mi caso", "Llamar ahora"]),
    ]
    for kicker_text, heading, paragraphs in ordered:
        h = doc.add_heading(heading, level=1)
        h.runs[0].font.color.rgb = RGBColor(11, 43, 107)
        k = doc.add_paragraph()
        rr = k.add_run(kicker_text.upper())
        rr.bold = True
        rr.font.color.rgb = RGBColor(224, 153, 0)
        for item in paragraphs:
            doc.add_paragraph(item)
    doc.save(path)


def validate(markup, preview):
    assert "[et_pb_code" not in markup
    assert "[et_pb_text" not in markup
    assert markup.count("[et_pb_section") == 12
    assert markup.count("[et_pb_contact_form") == 1
    assert markup.count("[et_pb_toggle") == 4
    for name in ["section", "row", "column", "blurb", "image", "button", "number_counter", "contact_form", "contact_field", "toggle"]:
        assert markup.count(f"[et_pb_{name}") == markup.count(f"[/et_pb_{name}]")
    bad = ["Ã", "Â", "â", "�", "&lt;p&gt;", "&gt;"]
    for token in bad:
        assert token not in markup
        assert token not in preview


def main():
    data = build_json()
    markup = next(iter(data["data"].values()))
    preview = build_html()
    validate(markup, preview)
    json_path = OUT / "01-Servicios-PORTAL-PREMIUM-DIVI4-NATIVO.json"
    html_path = OUT / "01-Servicios-PORTAL-PREMIUM-DIVI4-NATIVO-Preview.html"
    docx_path = OUT / "01-Servicios-PORTAL-PREMIUM-DIVI4-Textos.docx"
    json_path.write_text(json.dumps(data, ensure_ascii=False, separators=(",", ":")), encoding="utf-8")
    html_path.write_text(preview, encoding="utf-8")
    build_docx(docx_path)
    zip_path = OUT / "01-Servicios-PORTAL-PREMIUM-DIVI4-NATIVO.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        z.write(json_path, json_path.name)
        z.write(html_path, html_path.name)
        z.write(docx_path, docx_path.name)
    manifest = {
        "divi_version": V,
        "sections": 12,
        "code_modules": 0,
        "text_modules": 0,
        "native_contact_forms": 1,
        "native_toggles": 4,
        "files": [json_path.name, html_path.name, docx_path.name, zip_path.name],
    }
    (OUT / "VALIDACION.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
